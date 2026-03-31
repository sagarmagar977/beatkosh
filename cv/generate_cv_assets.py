from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parent
PDF_PATH = ROOT / "sagar-thapa-full-stack-cv.pdf"
RTF_PATH = ROOT / "sagar-thapa-full-stack-cv.rtf"
DOCX_PATH = ROOT / "sagar-thapa-full-stack-cv.docx"


CONTACT_LINES = [
    "Paskiot, Budhanilkantha",
    "+977-9804458369",
    "Sagarthapamagar977@gmail.com",
    "LinkedIn: https://www.linkedin.com/in/sagar-og-magar-656756210/",
    "Portfolio: https://sagarmagar977.github.io/MyPortfolio-/",
    "GitHub: https://github.com/sagarmagar977",
]

ABOUT = (
    "Full Stack Software Developer with hands-on experience building end-to-end web applications "
    "using Python, Django, Django REST Framework, React, Next.js, Tailwind CSS, and PostgreSQL/SQLite. "
    "Currently pursuing a Master of Computer Applications (MCA) and focused on developing scalable backend "
    "systems, modern frontend interfaces, secure authentication flows, and API-driven products. Built "
    "full-stack academic and personal projects with particular strength in backend architecture, REST APIs, "
    "role-based systems, and product-oriented problem solving."
)

EXPERIENCE = [
    (
        "2021-2023",
        "Math and Science Teacher (Class 8 - Class 10)",
        "Arunodaya Secondary School, Bhimad-2, Tanahun",
        [
            "Delivered mathematics and science instruction to secondary-level students.",
            "Strengthened communication, mentoring, and structured problem-solving skills now applied in software development.",
        ],
    ),
    (
        "2021-2022",
        "Math and English Instructor (Part Time)",
        "Ex-Army Fitness Training Center, Bhimad-6",
        [
            "Taught mathematics and English in a part-time instructional role.",
            "Built discipline in lesson planning, learner support, and clear communication.",
        ],
    ),
    (
        "2015-2017",
        "Math and Science Teacher (Class 6 - Class 8)",
        "Shree Bal Shaikshanik Secondary School, Bhimad-3, Tanahun",
        [],
    ),
]

EDUCATION = [
    ("Running", "Master of Computer Applications (MCA)", "Kantipur City College", "PUEB"),
    ("October 2023", "Post Graduate Diploma in Computer Applications (PGDCA) - CGPA 2.8", "Kantipur City College", "PUEB"),
    ("April 2021", "Bachelor of Science (B.Sc.) - 68.7%", "Sardar Patel University", "SPUEB"),
    ("December 2014", "+2 Science - 63.4%", "Sagarmatha Higher Secondary School", "HSEB"),
    ("June 2012", "SLC - 78.38%", "Amar Bhupu Secondary School", "NEB"),
]

TRAINING = [
    ("December 2023", "Python and Django (3 Months)", "Mindrisers Institute of Technology, Putalisadak"),
]

PROJECT_TITLE = "BeatKosh - Full Stack Music Marketplace and Freelance Collaboration Platform"
PROJECT_LINK = "https://github.com/sagarmagar977/beatkosh"
PROJECT_OVERVIEW = (
    "Designed and developed a full-stack platform for the Nepali music ecosystem that combines a "
    "multi-vendor beat marketplace with a freelance collaboration workflow for artists and producers. "
    "The product was planned and implemented as a backend-first system with a modern Next.js frontend, "
    "allowing users to browse and purchase beats, manage producer profiles, handle payments and wallet "
    "flows, and hire producers for milestone-based music projects."
)
PROJECT_STACK = (
    "Python, Django, Django REST Framework, React, Next.js, TypeScript, Tailwind CSS, "
    "PostgreSQL/SQLite, JWT, OpenAPI"
)
PROJECT_HIGHLIGHTS = [
    "Built backend modules for accounts, beats, catalog, orders, payments, projects, messaging, reviews, verification, analytics, and wallet workflows.",
    "Implemented role-aware user flows for artist and producer modes, including profile handling and permission-based access patterns.",
    "Developed marketplace APIs for beat listing, beat detail, trending content, bundles, beat tapes, and media upload workflows.",
    "Created collaboration features for project requests, proposals, milestones, deliverables, conversations, and messaging between artists and producers.",
    "Added payment and order flows with wallet credit handling, webhook-ready payment architecture, and download access rules.",
    "Built a Next.js frontend that connects to backend APIs for authentication, marketplace browsing, orders, projects, verification, and dashboard pages.",
]

SKILLS = [
    "Languages: Python, JavaScript, TypeScript, HTML, CSS, C, C++",
    "Backend: Django, Django REST Framework, JWT authentication, REST APIs, OpenAPI/Swagger",
    "Frontend: React, Next.js, Tailwind CSS, responsive UI development",
    "Database: PostgreSQL, SQLite, MySQL",
    "Tools: Git, GitHub, Postman, VS Code, npm, PowerShell",
]

LANGUAGES = ["English", "Nepali", "Hindi"]
REFERENCE = "Suru Basnet - Instructor, Mindrisers Institute of Technology (available upon request)"


def build_pdf() -> None:
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
        title="Sagar Thapa CV",
        author="OpenAI Codex",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=TA_LEFT,
        textColor=colors.black,
        spaceAfter=4,
    )
    contact_style = ParagraphStyle(
        "Contact",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=12,
        textColor=colors.black,
        spaceAfter=2,
    )
    section_style = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=13,
        textColor=colors.HexColor("#1f3a5f"),
        spaceBefore=8,
        spaceAfter=5,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        textColor=colors.black,
        spaceAfter=4,
    )
    small_bold = ParagraphStyle(
        "SmallBold",
        parent=body_style,
        fontName="Helvetica-Bold",
        spaceAfter=2,
    )

    story = []
    story.append(Paragraph("SAGAR THAPA", title_style))
    for line in CONTACT_LINES:
        story.append(Paragraph(line, contact_style))
    story.append(Spacer(1, 4))

    story.append(Paragraph("About Me", section_style))
    story.append(Paragraph(ABOUT, body_style))

    story.append(Paragraph("Experience", section_style))
    for date_text, role, org, bullets in EXPERIENCE:
        story.append(Paragraph(f"{date_text}  <b>{role}</b>", body_style))
        story.append(Paragraph(org, body_style))
        if bullets:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(item, body_style), leftIndent=8) for item in bullets],
                    bulletType="bullet",
                    bulletFontName="Helvetica",
                    bulletFontSize=7,
                    leftIndent=12,
                )
            )
        story.append(Spacer(1, 3))

    story.append(Paragraph("Education", section_style))
    for date_text, degree, school, board in EDUCATION:
        story.append(Paragraph(f"{date_text}  <b>{board}</b>", body_style))
        story.append(Paragraph(degree, body_style))
        story.append(Paragraph(school, body_style))
        story.append(Spacer(1, 2))

    story.append(Paragraph("Training/Certificates", section_style))
    for date_text, title, provider in TRAINING:
        story.append(Paragraph(f"{date_text}  <b>{title}</b>", body_style))
        story.append(Paragraph(provider, body_style))

    story.append(Paragraph("Projects", section_style))
    story.append(Paragraph(f"1. {PROJECT_TITLE}", small_bold))
    story.append(Paragraph(f"<b>Overview:</b> {PROJECT_OVERVIEW}", body_style))
    story.append(Paragraph(f"<b>Tech Used:</b> {PROJECT_STACK}", body_style))
    story.append(Paragraph(f"<b>Project Link:</b> {PROJECT_LINK}", body_style))
    story.append(Paragraph("<b>Key Highlights:</b>", body_style))
    story.append(
        ListFlowable(
            [ListItem(Paragraph(item, body_style), leftIndent=8) for item in PROJECT_HIGHLIGHTS],
            bulletType="bullet",
            bulletFontName="Helvetica",
            bulletFontSize=7,
            leftIndent=12,
        )
    )

    story.append(Paragraph("Skills", section_style))
    story.append(
        ListFlowable(
            [ListItem(Paragraph(item, body_style), leftIndent=8) for item in SKILLS],
            bulletType="bullet",
            bulletFontName="Helvetica",
            bulletFontSize=7,
            leftIndent=12,
        )
    )

    story.append(Paragraph("Languages", section_style))
    story.append(Paragraph(", ".join(LANGUAGES), body_style))

    story.append(Paragraph("References", section_style))
    story.append(Paragraph(REFERENCE, body_style))

    doc.build(story)


def rtf_par(text: str, bold: bool = False, size_half_points: int = 20, indent: int = 0) -> str:
    safe = escape(text).replace("\n", r"\line ")
    prefix = ""
    if bold:
        prefix += r"\b "
    if indent:
        prefix += rf"\li{indent} "
    suffix = r" \b0" if bold else ""
    return rf"\pard\sa120\fs{size_half_points} {prefix}{safe}{suffix}\par" + "\n"


def build_rtf() -> None:
    parts = [
        r"{\rtf1\ansi\deff0",
        r"{\fonttbl{\f0 Arial;}}",
        r"\viewkind4\uc1",
        rtf_par("SAGAR THAPA", bold=True, size_half_points=28),
    ]
    for line in CONTACT_LINES:
        parts.append(rtf_par(line, size_half_points=20))

    parts.append(rtf_par("About Me", bold=True, size_half_points=22))
    parts.append(rtf_par(ABOUT))

    parts.append(rtf_par("Experience", bold=True, size_half_points=22))
    for date_text, role, org, bullets in EXPERIENCE:
        parts.append(rtf_par(f"{date_text} - {role}", bold=True))
        parts.append(rtf_par(org))
        for bullet in bullets:
            parts.append(rtf_par(f"- {bullet}", indent=360))

    parts.append(rtf_par("Education", bold=True, size_half_points=22))
    for date_text, degree, school, board in EDUCATION:
        parts.append(rtf_par(f"{date_text} - {board}", bold=True))
        parts.append(rtf_par(degree))
        parts.append(rtf_par(school))

    parts.append(rtf_par("Training/Certificates", bold=True, size_half_points=22))
    for date_text, title, provider in TRAINING:
        parts.append(rtf_par(f"{date_text} - {title}", bold=True))
        parts.append(rtf_par(provider))

    parts.append(rtf_par("Projects", bold=True, size_half_points=22))
    parts.append(rtf_par(f"1. {PROJECT_TITLE}", bold=True))
    parts.append(rtf_par(f"Overview: {PROJECT_OVERVIEW}"))
    parts.append(rtf_par(f"Tech Used: {PROJECT_STACK}"))
    parts.append(rtf_par(f"Project Link: {PROJECT_LINK}"))
    parts.append(rtf_par("Key Highlights:", bold=True))
    for item in PROJECT_HIGHLIGHTS:
        parts.append(rtf_par(f"- {item}", indent=360))

    parts.append(rtf_par("Skills", bold=True, size_half_points=22))
    for item in SKILLS:
        parts.append(rtf_par(f"- {item}", indent=360))

    parts.append(rtf_par("Languages", bold=True, size_half_points=22))
    parts.append(rtf_par(", ".join(LANGUAGES)))

    parts.append(rtf_par("References", bold=True, size_half_points=22))
    parts.append(rtf_par(REFERENCE))
    parts.append("}")

    RTF_PATH.write_text("".join(parts), encoding="utf-8")


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = title.add_run("SAGAR THAPA")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)

    for line in CONTACT_LINES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Arial"
        r.font.size = Pt(9)

    add_docx_heading(doc, "About Me")
    add_docx_paragraph(doc, ABOUT)

    add_docx_heading(doc, "Experience")
    for date_text, role, org, bullets in EXPERIENCE:
        add_docx_paragraph(doc, f"{date_text} - {role}", bold=True)
        add_docx_paragraph(doc, org)
        for bullet in bullets:
            add_docx_bullet(doc, bullet)

    add_docx_heading(doc, "Education")
    for date_text, degree, school, board in EDUCATION:
        add_docx_paragraph(doc, f"{date_text} - {board}", bold=True)
        add_docx_paragraph(doc, degree)
        add_docx_paragraph(doc, school)

    add_docx_heading(doc, "Training/Certificates")
    for date_text, title, provider in TRAINING:
        add_docx_paragraph(doc, f"{date_text} - {title}", bold=True)
        add_docx_paragraph(doc, provider)

    add_docx_heading(doc, "Projects")
    add_docx_paragraph(doc, f"1. {PROJECT_TITLE}", bold=True)
    add_docx_paragraph(doc, f"Overview: {PROJECT_OVERVIEW}")
    add_docx_paragraph(doc, f"Tech Used: {PROJECT_STACK}")
    add_docx_paragraph(doc, f"Project Link: {PROJECT_LINK}")
    add_docx_paragraph(doc, "Key Highlights:", bold=True)
    for item in PROJECT_HIGHLIGHTS:
        add_docx_bullet(doc, item)

    add_docx_heading(doc, "Skills")
    for item in SKILLS:
        add_docx_bullet(doc, item)

    add_docx_heading(doc, "Languages")
    add_docx_paragraph(doc, ", ".join(LANGUAGES))

    add_docx_heading(doc, "References")
    add_docx_paragraph(doc, REFERENCE)

    doc.save(str(DOCX_PATH))


def add_docx_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_docx_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)


def add_docx_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
if __name__ == "__main__":
    build_pdf()
    build_rtf()
    build_docx()
    print(f"Created: {PDF_PATH.name}")
    print(f"Created: {RTF_PATH.name}")
    print(f"Created: {DOCX_PATH.name}")
